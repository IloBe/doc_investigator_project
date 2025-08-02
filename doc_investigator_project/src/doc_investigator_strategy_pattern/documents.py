# src/doc_investigator_strategy_pattern/documents.py

"""
Document processing module for the Document Investigator application.

Implements the Strategy design pattern to handle text extraction from
various file types.

Classes:
    - InvalidFileTypeException: Custom exception for unsupported file types
    - DocumentLoaderStrategy: Abstract base class for all file loaders
    - PDFLoaderStrategy, DocxLoaderStrategy, etc.: Concrete strategy implementations
    - DocumentProcessor: Context class that uses a strategy to process files
"""

# ----------
# Imports
# ----------
import abc
import os
from typing import Any, Dict, List, Iterator, Generator

import docx
import fitz  # PyMuPDF
import openpyxl
from loguru import logger

# ----------
# Coding
# ----------

# --- Custom Exception ---
class InvalidFileTypeException(Exception):
    """Custom exception raised for unsupported file types."""
    pass

# --- Strategy Pattern for Document Loading ---
class DocumentLoaderStrategy(abc.ABC):
    """Abstract base class for a document loading strategy."""

    @abc.abstractmethod
    def load(self, file_path: str) -> str:
        """Loads a file and returns its text content as a string."""
        pass

class PDFLoaderStrategy(DocumentLoaderStrategy):
    """Strategy for loading text from PDF files."""
    def load(self, file_path: str) -> str:
        try:
            text = ""
            with fitz.open(file_path) as doc:
                content = "\n".join(page.get_text() for page in doc)
                
            logger.debug(f"Successfully extracted text from PDF: {os.path.basename(file_path)}")
            return content
        except fitz.fitz.PyMuPDFError as e:  # specific error class for PyMuPDF (corrupt files, password protection, ...)
            logger.error(f"Error loading PDF: '{file_path}': {e}", exc_info = True)
            return f"[Error processing PDF: {os.path.basename(file_path)} - The file may be corrupt, password-protected or unreadable.]"
        except (FileNotFoundError) as e:
            logger.error(f"Error PDF file not found '{file_path}': {e}", exc_info = True)
            return f"[Error PDF file not found: {os.path.basename(file_path)}.]"

class DocxLoaderStrategy(DocumentLoaderStrategy):
    """Strategy for loading text from DOCX files, extracting paragraphs and tables."""
    
    def _iter_text(self, doc) -> Generator[str, None, None]:
        """Generator to yield all text blocks from paragraphs and tables."""
        # paragraphs
        for para in doc.paragraphs:
            yield para.text   
        # text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # join cell text with space or tab
                    yield cell.text
    
    def load(self, file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            content = "\n".join(self._iter_text(doc))
            logger.debug(f"Successfully extracted text from paragraphs and tables in DOCX: {os.path.basename(file_path)}")
            return content
            
        except (FileNotFoundError, docx.opc.exceptions.PackageNotFoundError):
            logger.error(f"File not found or is not a valid DOCX: '{file_path}'")
            return f"[Error: File not found or is not a valid DOCX file {os.path.basename(file_path)}]"
        except Exception as e:
            logger.error(f"Error loading DOCX '{file_path}': {e}", exc_info = True)
            return f"[Error processing DOCX: {os.path.basename(file_path)} - The file may be corrupt or incompatible.]"

class TextLoaderStrategy(DocumentLoaderStrategy):
    """Strategy for loading text from TXT files line by line, returns a single string whith whole content by success."""
    def load(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding = 'utf-8', errors = 'ignore') as f:
                # consumes file's iterator and joins all lines
                content = "".join(f)
            
            logger.debug(f"Successful extracted all lines from TXT: {os.path.basename(file_path)}")
            return content
        except FileNotFoundError:
            logger.error(f"File not found: '{file_path}'")
            return f"[Error: File not found: {os.path.basename(file_path)}]"
        except IOError as e:
            logger.error(f"Error reading TXT '{file_path}': {e}", exc_info = True)
            return f"[Error read processing TXT: {os.path.basename(file_path)}]"
        except Exception as e:
            logger.error(f"An unexpected error occurred with '{file_path}': {e}", exc_info = True)
            return f"[Error: An unexpected issue occurred with {os.path.basename(file_path)}]"

class ExcelLoaderStrategy(DocumentLoaderStrategy):
    """Strategy for loading text from XLSX files resp. its sheets."""
    def load(self, file_path: str) -> str:
        try:
            workbook = openpyxl.load_workbook(file_path, read_only = True)

            # generator usage for sheets and cells; loads one row at a time for constant O(1)
            def content_generator():
                for sheet_name in workbook.sheetnames:
                    yield f"--- Sheet: {sheet_name} ---"
                    sheet = workbook[sheet_name]
                    for row in sheet.iter_rows():
                        row_text = "\t".join(
                            str(cell.value) if cell.value is not None else ""
                            for cell in row
                        )
                        yield row_text
        
            # generator consumed, building final string
            logger.debug(f"Successfully extracted text from XLSX: {os.path.basename(file_path)}")
            return "\n".join(content_generator())
            
        except FileNotFoundError as e: # openpyxl can raise various errors
            logger.error(f"Error loading XLSX '{file_path}': {e}", exc_info = True)
            return f"Error processing XLSX: {os.path.basename(file_path)} - File may be corrupt."
        except Exception as e:
            logger.error(f"Failed to extract text from XLSX {os.path.basename(file_path)}: {e}")
            # Depending on requirements, you might want to return "" or re-raise
            return f"Failed to extract text from XLSX {os.path.basename(file_path)}: {e}"
        finally:
            # workbook closing to release file handle
            if 'workbook' in locals() and workbook is not None:
                workbook.close() # read_only mode: openpyxl requires explicit closing
                logger.debug(f"Successfully extracted text and closed workbook: {os.path.basename(file_path)}")


# --- Main Processor (Context Class) ---
class DocumentProcessor:
    """
    Context class uses the strategy to validate and process documents.

    Initialized with a set of supported file types and uses a dictionary
    of loader strategies to perform the actual text extraction.
    """
    def __init__(self, supported_extensions: List[str]):
        """
        Initializes the DocumentProcessor with a mapping of extensions to loaders.

        Args:
            supported_extensions (List[str]): A list of supported file extensions
                                              (e.g., ['.pdf', '.docx'])
        """
        self.supported_extensions: List[str] = supported_extensions
        self._strategies: Dict[str, DocumentLoaderStrategy] = {
            '.pdf': PDFLoaderStrategy(),
            '.docx': DocxLoaderStrategy(),
            '.txt': TextLoaderStrategy(),
            '.xlsx': ExcelLoaderStrategy(),
        }
        logger.info(f"DocumentProcessor initialized for types: {', '.join(supported_extensions)}")

    def validate_files(self, files: List[Any]) -> None:
        """
        Validates that all uploaded files have a supported extension.

        Args:
            files (List[Any]): A list of file-like objects (from Gradio)

        Raises:
            InvalidFileTypeException: If any file has an unsupported extension
        """
        logger.info(f"Validating {len(files)} uploaded files...")
        for file in files:
            file_name = os.path.basename(file.name)
            _, file_extension = os.path.splitext(file_name)
            if file_extension.lower() not in self.supported_extensions:
                error_msg = f"Unsupported file type: '{file_extension}' in file '{file_name}'."
                logger.warning(error_msg)
                raise InvalidFileTypeException(error_msg)
        logger.success(f"All {len(files)} files passed validation.")

    def process_files(self, files: List[Any]) -> str:
        """
        Extracts text from a list of pre-validated Gradio file objects.

        This method iterates through the files, selects the appropriate loading
        strategy based on file extension, and aggregates the content.

        Args:
            files (List[Any]): list of validated file-like objects

        Returns:
            str: combined text content of all files, with headers indicating source
        """
        logger.info(f"Processing {len(files)} validated files to extract text...")
        all_texts: List[str] = []
        for file in files:
            file_path: str = file.name
            file_name: str = os.path.basename(file_path)
            _, file_extension = os.path.splitext(file_path)

            strategy = self._strategies.get(file_extension.lower())
            # safeguard check, though validation should prevent this behaviour
            if not strategy:
                logger.warning(f"No loader strategy found for validated file '{file_name}'. Skipping.")
                continue

            text = strategy.load(file_path)
            all_texts.append(f"--- CONTENT FROM {file_name} ---\n{text}")

        logger.success("Text extraction from all files complete.")
        return "\n\n".join(all_texts)